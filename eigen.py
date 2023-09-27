
import numpy as np
import os
import cv2
from docx import Document
from docx.shared import Inches
from sklearn.decomposition import TruncatedSVD

data_path = r'C:\Users\Costin\Documents\optimizationtechS2\data\ORL\s1'
image_paths = [os.path.join(data_path, f) for f in os.listdir(data_path) if f.endswith('.pgm')]
images = [cv2.imread(f, cv2.IMREAD_GRAYSCALE) for f in image_paths]

print("Numar de poze incarcate :", len(images))

X = np.asarray(images)

n_samples = X.shape[0]
X = X.reshape(n_samples, -1)

mean_face = np.mean(X, axis=0)

X = X - mean_face

L = np.dot(X, X.T)

k = 10
svd = TruncatedSVD(n_components=k)
eigenvectors = svd.fit_transform(L)

eigenfaces = np.dot(X.T, eigenvectors)

eigenfaces_norm = np.linalg.norm(eigenfaces, axis=0)
eigenfaces = eigenfaces / eigenfaces_norm

for i in range(k):
    
    eigenface_1d = eigenfaces[:, i].flatten()
    
 
    eigenface_resized = cv2.resize(eigenface_1d, (64, 64))
    
    
    cv2.imshow("Eigenface {}".format(i), eigenface_resized)
    cv2.waitKey(0)



doc = Document()

doc.add_heading('Eigenfaces', level=0)

doc.add_paragraph(f'Savele Costin')


for i in range(k):
    
    eigenface_1d = eigenfaces[:, i].flatten()
    
    
    eigenface_resized = cv2.resize(eigenface_1d, (64, 64))
    
    
    cv2.imwrite("eigenface_{}.png".format(i), eigenface_resized)
    
    
    doc.add_paragraph('Eigenface {}'.format(i))
    doc.add_picture('eigenface_{}.png'.format(i), width=Inches(2))
    doc.add_paragraph()

doc.save('eigenfaces.docx')

cv2.destroyAllWindows()

